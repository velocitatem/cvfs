"""Generate the static demo CV DOCX used by DEMO mode in the webapp."""
from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_bullet(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Bullet")


def build(path: Path) -> None:
    doc = Document()

    # Name / contact
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = name_para.add_run("Alex Rivera")
    run.bold = True
    run.font.size = Pt(18)

    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.add_run("alex.rivera@email.com  ·  linkedin.com/in/alexrivera  ·  github.com/alexrivera")

    doc.add_paragraph()  # spacer

    # Summary
    add_heading(doc, "Summary", level=2)
    doc.add_paragraph(
        "Software engineer with 5 years of experience building distributed systems and "
        "machine learning pipelines at scale. Strong background in Python, Go, and cloud-native "
        "architectures. Passionate about developer tooling and open-source contribution."
    )

    # Experience
    add_heading(doc, "Experience", level=2)

    add_heading(doc, "Senior Software Engineer — Acme Corp", level=3)
    doc.add_paragraph("Jan 2022 – Present  ·  San Francisco, CA")
    add_bullet(doc, "Led migration of monolithic data pipeline to distributed microservices, reducing p99 latency by 40%.")
    add_bullet(doc, "Designed and shipped an internal feature flag system used by 50+ engineers across 3 teams.")
    add_bullet(doc, "Mentored 4 junior engineers and ran weekly technical design review sessions.")

    add_heading(doc, "Software Engineer — DataFlow Inc", level=3)
    doc.add_paragraph("Aug 2019 – Dec 2021  ·  Remote")
    add_bullet(doc, "Built real-time streaming ingestion system processing 2M events/day using Kafka and Flink.")
    add_bullet(doc, "Developed Python SDK for internal data platform, adopted by 8 product teams.")
    add_bullet(doc, "Contributed PyTorch-based anomaly detection model achieving 92% precision on production traffic.")

    # Education
    add_heading(doc, "Education", level=2)
    add_heading(doc, "B.S. Computer Science — State University", level=3)
    doc.add_paragraph("Graduated May 2019  ·  GPA 3.8 / 4.0")
    add_bullet(doc, "Senior thesis: Efficient approximate nearest-neighbour search for high-dimensional embeddings.")

    # Skills
    add_heading(doc, "Skills", level=2)
    skills_para = doc.add_paragraph()
    skills_para.add_run("Languages: ").bold = True
    skills_para.add_run("Python, Go, TypeScript, SQL")

    infra_para = doc.add_paragraph()
    infra_para.add_run("Infrastructure: ").bold = True
    infra_para.add_run("Kubernetes, AWS, GCP, Terraform, Docker")

    ml_para = doc.add_paragraph()
    ml_para.add_run("ML / Data: ").bold = True
    ml_para.add_run("PyTorch, scikit-learn, Spark, Kafka, dbt")

    doc.save(path)
    print(f"Saved demo CV to {path}")


if __name__ == "__main__":
    out = Path(sys.argv[1]) if len(sys.argv) > 1 else Path("apps/webapp/public/demo-cv.docx")
    out.parent.mkdir(parents=True, exist_ok=True)
    build(out)
