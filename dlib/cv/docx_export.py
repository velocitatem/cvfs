from __future__ import annotations

from collections import defaultdict
from io import BytesIO

from docx import Document

from .parser import _detect_block_type


def _path_to_para_map(doc: Document) -> dict[str, int]:
    counters: defaultdict[str, int] = defaultdict(int)
    result: dict[str, int] = {}
    for idx, para in enumerate(doc.paragraphs):
        if not para.text.strip():
            continue
        block_type = _detect_block_type(getattr(para.style, "name", None), para)
        counters[block_type] += 1
        result[f"{block_type}[{counters[block_type]}]"] = idx
    return result


def _replace_para_text(para, new_text: str) -> None:
    """Replace paragraph text preserving the first run's character formatting."""
    if not para.runs:
        para.add_run(new_text)
        return
    first = para.runs[0]
    for run in para.runs[1:]:
        run.text = ""
    first.text = new_text


def _remove_paragraph(paragraph) -> None:
    p = paragraph._element
    p.getparent().remove(p)


def generate_patched_docx(
    original_bytes: bytes, structured_blocks: list[dict]
) -> bytes:
    """Return DOCX bytes with text patches from structured_blocks applied.

    Compares each block's text against the original paragraph and replaces it
    when different. Blocks absent from structured_blocks are removed.
    """
    if not structured_blocks:
        return original_bytes

    doc = Document(BytesIO(original_bytes))
    path_map = _path_to_para_map(doc)

    original_paths = set(path_map.keys())
    patched = {b["path"]: b["text"] for b in structured_blocks}
    patched_paths = set(patched.keys())

    # Apply text replacements first (indices stay stable)
    for path, new_text in patched.items():
        idx = path_map.get(path)
        if idx is None:
            continue
        para = doc.paragraphs[idx]
        if para.text.strip() != new_text:
            _replace_para_text(para, new_text)

    # Remove blocks no longer present; process in reverse index order
    removed = sorted(
        [path_map[p] for p in (original_paths - patched_paths) if p in path_map],
        reverse=True,
    )
    for idx in removed:
        _remove_paragraph(doc.paragraphs[idx])

    out = BytesIO()
    doc.save(out)
    return out.getvalue()
