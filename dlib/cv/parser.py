from __future__ import annotations

from collections import defaultdict
from io import BytesIO
from typing import Iterable

from docx import Document

from .schema import StructuredBlock, StructuredDocument


def _detect_block_type(style_name: str | None, paragraph) -> str:
    style = (style_name or "").lower()
    if style.startswith("heading"):
        return "heading"
    if (
        "bullet" in style
        or "list" in style
        or getattr(paragraph, "style", None)
        and getattr(paragraph.style, "name", "").lower().startswith("list")
    ):
        return "bullet"
    return "text"


def _build_path(block_type: str, counter: int, extra: str | None = None) -> str:
    suffix = f"{block_type}[{counter}]"
    if extra:
        return f"{suffix}.{extra}"
    return suffix


def parse_docx_bytes(
    file_bytes: bytes, *, version_label: str | None = None
) -> StructuredDocument:
    document = Document(BytesIO(file_bytes))
    counters: defaultdict[str, int] = defaultdict(int)
    blocks: list[StructuredBlock] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        block_type = _detect_block_type(
            getattr(paragraph.style, "name", None), paragraph
        )
        counters[block_type] += 1
        keywords = summarize_keywords([text])
        blocks.append(
            StructuredBlock(
                path=_build_path(block_type, counters[block_type]),
                block_type="heading"
                if block_type == "heading"
                else ("bullet" if block_type == "bullet" else "text"),
                text=text,
                keywords=keywords,
                metadata={
                    "style": getattr(getattr(paragraph, "style", None), "name", "")
                },
            )
        )

    for table_index, table in enumerate(document.tables):
        for row_index, row in enumerate(table.rows):
            for cell_index, cell in enumerate(row.cells):
                text = cell.text.strip()
                if not text:
                    continue
                counters["table"] += 1
                blocks.append(
                    StructuredBlock(
                        path=_build_path(
                            "table",
                            counters["table"],
                            extra=f"{row_index}-{cell_index}",
                        ),
                        block_type="table",
                        text=text,
                        keywords=summarize_keywords([text]),
                        metadata={
                            "table_index": table_index,
                            "row": row_index,
                            "cell": cell_index,
                        },
                    )
                )

    return StructuredDocument(version_label=version_label, blocks=blocks)


def summarize_keywords(lines: Iterable[str], *, max_keywords: int = 6) -> list[str]:
    terms: dict[str, int] = {}
    for line in lines:
        for raw in line.split():
            cleaned = raw.strip().strip(",.;:()[]").lower()
            if len(cleaned) <= 2:
                continue
            terms[cleaned] = terms.get(cleaned, 0) + 1
    return [
        term
        for term, _ in sorted(terms.items(), key=lambda kv: kv[1], reverse=True)[
            :max_keywords
        ]
    ]
